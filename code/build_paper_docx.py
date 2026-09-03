from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "paper" / "main.md"
OUT = ROOT / "paper" / "final" / "24C_农作物种植策略论文.docx"
EQUATION_DIR = ROOT / "paper" / "equations"

# narrative_proposal preset, with named academic-paper overrides:
# A4 paper; 2.2 cm horizontal / 2.0 cm vertical margins; Chinese fonts;
# compact first-page title instead of a separate editorial cover.
PAGE_W_CM = 21.0
PAGE_H_CM = 29.7
MARGIN_X_CM = 2.2
MARGIN_Y_CM = 2.0
CONTENT_W_DXA = 9411
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F4F6F9"
MUTED = "666666"


def set_run_font(run, east="宋体", latin="Times New Roman", size=10.5,
                 bold=None, italic=None, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, east="宋体", latin="Times New Roman", size=9, color=MUTED)


def add_bottom_border(paragraph, color="D9E2F3", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.first_line_indent = Cm(0.74)

    for style_name, size, color, before, after, east in [
        ("Heading 1", 16, BLUE, 18, 10, "黑体"),
        ("Heading 2", 13, BLUE, 12, 6, "黑体"),
        ("Heading 3", 12, DARK_BLUE, 8, 4, "黑体"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        p = style.paragraph_format
        p.space_before = Pt(before)
        p.space_after = Pt(after)
        p.line_spacing = 1.0
        p.keep_with_next = True
        p.first_line_indent = Cm(0)

    for style_name in ["List Number", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        p = style.paragraph_format
        p.left_indent = Cm(0.95)
        p.first_line_indent = Cm(-0.49)
        p.space_after = Pt(4)
        p.line_spacing = 1.208


def configure_section(section):
    section.page_width = Cm(PAGE_W_CM)
    section.page_height = Cm(PAGE_H_CM)
    section.top_margin = Cm(MARGIN_Y_CM)
    section.bottom_margin = Cm(MARGIN_Y_CM)
    section.left_margin = Cm(MARGIN_X_CM)
    section.right_margin = Cm(MARGIN_X_CM)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("2024年全国大学生数学建模竞赛 · C题")
    set_run_font(run, east="宋体", latin="Times New Roman", size=8.5, color=MUTED)
    add_bottom_border(hp)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("— ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(fp)
    r = fp.add_run(" —")
    set_run_font(r, size=9, color=MUTED)


def latex_to_text(expr):
    expr = re.sub(r"\\tag\{[^}]+\}", "", expr)
    replacements = {
        r"\omega": "ω", r"\Omega": "Ω", r"\Pi": "Π", r"\pi": "π",
        r"\alpha": "α", r"\eta": "η", r"\xi": "ξ", r"\lambda": "λ",
        r"\tau": "τ", r"\sum": "∑", r"\le": "≤", r"\ge": "≥",
        r"\in": "∈", r"\max": "max", r"\min": "min", r"\left": "",
        r"\right": "", r"\qquad": "    ", r"\,": " ", r"\ ": " ",
        r"\begin{bmatrix}": "[", r"\end{bmatrix}": "]", r"\\": "; ",
    }
    for old, new in replacements.items():
        expr = expr.replace(old, new)
    expr = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", expr)
    expr = re.sub(r"\\text\{([^{}]+)\}", r"\1", expr)
    expr = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", expr)
    expr = expr.replace("&", " ")
    expr = expr.replace("{", "").replace("}", "")
    expr = re.sub(r"\s+", " ", expr).strip()
    return expr


INLINE_RE = re.compile(r"(\*\*.*?\*\*|\$.*?\$)")


def add_inline(paragraph, text, *, base_size=10.5, bold_default=False):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, bold=bold_default)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, east="黑体", size=base_size, bold=True)
        else:
            run = paragraph.add_run(latex_to_text(token[1:-1]))
            set_run_font(run, east="Cambria Math", latin="Cambria Math", size=base_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, bold=bold_default)


FORMULA_TEXT = {
    "1": "Qⱼₜₛ^ω = ∑ᵢ qⱼ,Lᵢ,s^ω xᵢⱼₜₛ",
    "2": "0 ≤ uⱼₜₛ^ω ≤ Qⱼₜₛ^ω,     uⱼₜₛ^ω ≤ dⱼₛ^t,ω",
    "3": "Π^ω = ∑ⱼ,ₜ,ₛ pⱼ^t,ω [uⱼₜₛ^ω + r eⱼₜₛ^ω] − ∑ᵢ,ⱼ,ₜ,ₛ cⱼ,Lᵢ,s^t,ω xᵢⱼₜₛ",
    "4": "∑ⱼ xᵢⱼₜₛ ≤ Aᵢ",
    "5": "mᵢ yᵢⱼₜₛ ≤ xᵢⱼₜₛ ≤ Aᵢ yᵢⱼₜₛ",
    "6": "∑τ₌ₜ^{t+2} ∑ⱼ∈Jbean,ₛ xᵢⱼτₛ ≥ Aᵢ",
    "7": "max  Π − 0.50∑ⱼ,ₜ,ₛ pⱼ eⱼₜₛ − λ∑ᵢ,ⱼ,ₜ,ₛ yᵢⱼₜₛ",
    "8": "ξω ≥ Lω − η,     ξω ≥ 0",
    "9": "−(1/|Ω|)∑ω Πω + 0.15[η + 1/((1−0.90)|Ω|)∑ω ξω] + λ∑ yᵢⱼₜₛ",
    "10": "R = [ 1       −0.35    0.25     0.30\n"
          "     −0.35     1      −0.10    −0.20\n"
          "      0.25    −0.10    1        0.40\n"
          "      0.30    −0.20    0.40     1    ]",
}


def render_formula(tag, text):
    EQUATION_DIR.mkdir(parents=True, exist_ok=True)
    path = EQUATION_DIR / f"eq_{tag}.png"
    font = ImageFont.truetype(r"C:\Windows\Fonts\cambria.ttc", 42)
    probe = Image.new("RGB", (20, 20), "white")
    draw = ImageDraw.Draw(probe)
    box = draw.multiline_textbbox((0, 0), text, font=font, spacing=7)
    width = box[2] - box[0] + 36
    height = box[3] - box[1] + 28
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.multiline_text((18, 10 - box[1]), text, font=font, fill="black", spacing=7,
                        align="left")
    image.save(path, dpi=(288, 288))
    return path, width


def add_equation(doc, raw):
    tag_match = re.search(r"\\tag\{([^}]+)\}", raw)
    tag = f"({tag_match.group(1)})" if tag_match else ""
    tag_value = tag_match.group(1) if tag_match else "x"
    formula = FORMULA_TEXT.get(tag_value, latex_to_text(raw))
    image_path, pixel_width = render_formula(tag_value, formula)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [8210, 1201])
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "nil")
            borders.append(node)
        tc_pr.append(borders)
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    target_cm = min(13.9, max(3.5, pixel_width / 288 * 2.54))
    run.add_picture(str(image_path), width=Cm(target_cm))
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(tag)
    set_run_font(run, east="Cambria Math", latin="Cambria Math", size=10.2)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_markdown_table(doc, rows):
    cols = max(len(r) for r in rows)
    widths = [CONTENT_W_DXA // cols] * cols
    widths[-1] += CONTENT_W_DXA - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                shade_cell(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, row[c_idx] if c_idx < len(row) else "", base_size=9.0,
                       bold_default=(r_idx == 0))
    set_repeat_table_header(table.rows[0])


def resolve_image(link):
    name = Path(link).name
    candidate = ROOT / "paper" / "figures" / name
    return candidate


def add_figure(doc, alt, link):
    path = resolve_image(link)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.8))


def add_title_block(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, east="黑体", latin="Times New Roman", size=20, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("2024年高教社杯全国大学生数学建模竞赛 · C题")
    set_run_font(r, east="宋体", latin="Times New Roman", size=10, color=MUTED)
    add_bottom_border(p, color=BLUE, size="10")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0].lstrip("# ").strip()
    add_title_block(doc, title)

    i = 1
    in_equation = False
    equation_parts = []
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if in_equation:
            equation_parts.append(line)
            if line.endswith("$$"):
                expr = " ".join(equation_parts)
                expr = expr[2:-2].strip()
                add_equation(doc, expr)
                in_equation = False
                equation_parts = []
            i += 1
            continue
        if line.startswith("$$"):
            if line.endswith("$$") and len(line) > 4:
                add_equation(doc, line[2:-2].strip())
            else:
                in_equation = True
                equation_parts = [line]
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        image_match = re.fullmatch(r"!\[(.*)\]\((.*)\)", line)
        if image_match:
            add_figure(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:], base_size=13, bold_default=True)
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:], base_size=16, bold_default=True)
        elif line.startswith("# "):
            pass
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s*", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, text)
        else:
            p = doc.add_paragraph()
            if line.startswith("**图"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(8)
                add_inline(p, line, base_size=9.5)
            elif line.startswith("**关键词：**"):
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(10)
                add_inline(p, line)
            elif line.startswith("[") and re.match(r"^\[\d+\]", line):
                p.paragraph_format.first_line_indent = Cm(-0.74)
                p.paragraph_format.left_indent = Cm(0.74)
                p.paragraph_format.space_after = Pt(4)
                add_inline(p, line, base_size=9.5)
            else:
                add_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = title
    props.subject = "2024年全国大学生数学建模竞赛C题：农作物的种植策略"
    props.author = ""
    props.keywords = "种植规划, 混合整数线性规划, 滚动优化, 情景分析, CVaR"
    props.comments = "Generated from verified modeling artifacts."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
